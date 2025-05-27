from typing import Dict, List, Optional, Union, Any
import os
from pathlib import Path
import magic
from pypdf import PdfReader
import docx
import pdfplumber
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
import boto3
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import pickle
import base64
from email.mime.text import MIMEText
import json
import hashlib
import redis
import cv2
import numpy as np
from PIL import Image
import ezdxf
import io
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class LLMCache:
    """Cache for LLM responses to improve performance"""
    def __init__(self, redis_url: Optional[str] = None):
        self.redis_client = None
        if redis_url:
            try:
                self.redis_client = redis.from_url(redis_url)
                logger.info("Connected to Redis cache")
            except Exception as e:
                logger.warning(f"Failed to connect to Redis: {e}")
                self.redis_client = None

    def get(self, key: str) -> Optional[str]:
        """Get value from cache"""
        if not self.redis_client:
            return None
        try:
            return self.redis_client.get(key)
        except Exception as e:
            logger.warning(f"Cache get error: {e}")
            return None

    def set(self, key: str, value: str, expire: int = 3600) -> bool:
        """Set value in cache with expiration"""
        if not self.redis_client:
            return False
        try:
            return self.redis_client.set(key, value, ex=expire)
        except Exception as e:
            logger.warning(f"Cache set error: {e}")
            return False

class DocumentProcessor:
    """Base class for document processing"""
    def __init__(self, redis_url: Optional[str] = None):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        self.llm = ChatOpenAI(
            model_name="gpt-3.5-turbo",
            temperature=0
        )
        self.cache = LLMCache(redis_url)
        self.metadata_prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert at analyzing construction and security system requirements.
            Extract the following information from the provided text:
            - Project type (e.g., commercial, residential, industrial, healthcare, educational)
            - Location (city, state, country)
            - Square footage (if mentioned)
            - Number of floors (if mentioned)
            - Security requirements (if any)
            - Fire safety requirements (if any)
            - Building codes and standards mentioned (e.g., NFPA, UL, NEC)
            - Special requirements (e.g., ADA compliance, LEED certification)
            - Timeline and deadlines
            - Budget constraints (if mentioned)
            - Existing systems to be integrated with
            - Environmental considerations
            
            Return the information in JSON format with these exact keys:
            {
                "project_type": string or null,
                "location": string or null,
                "square_footage": number or null,
                "number_of_floors": number or null,
                "security_requirements": list of strings or null,
                "fire_safety_requirements": list of strings or null,
                "building_codes": list of strings or null,
                "special_requirements": list of strings or null,
                "timeline": string or null,
                "budget_constraints": string or null,
                "existing_systems": list of strings or null,
                "environmental_considerations": list of strings or null
            }
            
            If any information is not found, use null for that field."""),
            ("user", "{text}")
        ])

    def _get_cache_key(self, content: str) -> str:
        """Generate a cache key for the content"""
        return hashlib.md5(content.encode()).hexdigest()

    def extract_metadata(self, content: str) -> Dict:
        """Extract key metadata from document content using LLM with caching"""
        try:
            # Get the first 4000 characters for metadata extraction
            truncated_content = content[:4000]
            cache_key = self._get_cache_key(truncated_content)
            
            # Try to get from cache
            cached_result = self.cache.get(cache_key)
            if cached_result:
                logger.info("Using cached metadata")
                return json.loads(cached_result)
            
            # Create the chain
            chain = self.metadata_prompt | self.llm
            
            # Get the response
            response = chain.invoke({"text": truncated_content})
            
            # Parse the JSON response
            metadata = json.loads(response.content)
            
            # Cache the result
            self.cache.set(cache_key, response.content)
            
            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            # Return default metadata if extraction fails
            return {
                "project_type": None,
                "location": None,
                "square_footage": None,
                "number_of_floors": None,
                "security_requirements": None,
                "fire_safety_requirements": None,
                "building_codes": None,
                "special_requirements": None,
                "timeline": None,
                "budget_constraints": None,
                "existing_systems": None,
                "environmental_considerations": None
            }

    def process(self, content: str) -> List[Document]:
        """Process document content into chunks"""
        chunks = self.text_splitter.split_text(content)
        metadata = self.extract_metadata(content)
        return [Document(page_content=chunk, metadata=metadata) for chunk in chunks]

class PDFProcessor(DocumentProcessor):
    """Process PDF documents"""
    def process_file(self, file_path: Union[str, Path]) -> List[Document]:
        content = ""
        # Try pypdf first
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                content += page.extract_text() or ""
        except Exception:
            # Fallback to pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    content += page.extract_text() or ""
        
        return self.process(content)

class DocxProcessor(DocumentProcessor):
    """Process DOCX documents"""
    def process_file(self, file_path: Union[str, Path]) -> List[Document]:
        doc = docx.Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return self.process(content)

class ImageProcessor(DocumentProcessor):
    """Process image files using OpenAI's Vision API with specialized analysis for fire alarms and security systems"""
    def __init__(self, redis_url: Optional[str] = None):
        super().__init__(redis_url)
        self.vision_model = ChatOpenAI(
            model_name="gpt-4-vision-preview",
            temperature=0,
            max_tokens=2000
        )
        self.vision_prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert at analyzing fire alarm and security system technical drawings, blueprints, and schematics.
            Your expertise includes NFPA, NEC, and UL standards compliance.
            
            Analyze the provided image and extract the following information:
            
            1. Document Classification:
            - Type (e.g., floor plan, riser diagram, single-line diagram, device layout)
            - Scale and units
            - Drawing version/revision
            - Date of drawing
            
            2. Fire Alarm System Components:
            - Control panels and their locations
            - Initiating devices (smoke detectors, heat detectors, pull stations)
            - Notification appliances (horns, strobes, speakers)
            - Fire alarm circuits and their types (SLC, NAC, etc.)
            - Battery calculations and power supplies
            - Emergency communication systems
            - NFPA 72 compliance indicators
            
            3. Security System Components:
            - Access control panels and readers
            - CCTV cameras and their types (dome, bullet, PTZ)
            - Intrusion detection devices
            - Intercom systems
            - Security zones and partitions
            - Network infrastructure (PoE switches, routers)
            
            4. Building Infrastructure:
            - Room numbers and names
            - Ceiling heights and types
            - Wall types and ratings
            - Electrical room locations
            - Network closet locations
            - Cable pathways and conduits
            
            5. Compliance and Standards:
            - NFPA code references
            - UL listed equipment
            - ADA compliance requirements
            - Local code requirements
            - Emergency egress paths
            
            6. Technical Specifications:
            - Device mounting heights
            - Coverage areas and spacing
            - Cable types and specifications
            - Power requirements
            - Network bandwidth requirements
            
            7. Integration Points:
            - Building management systems
            - Elevator controls
            - HVAC systems
            - Door hardware
            - Emergency power systems
            
            Return the information in JSON format with these exact keys:
            {
                "document_info": {
                    "type": string or null,
                    "scale": string or null,
                    "version": string or null,
                    "date": string or null
                },
                "fire_alarm_system": {
                    "control_panels": list of strings or null,
                    "initiating_devices": list of strings or null,
                    "notification_appliances": list of strings or null,
                    "circuits": list of strings or null,
                    "power_supplies": list of strings or null,
                    "emergency_communication": list of strings or null,
                    "nfpa_compliance": list of strings or null
                },
                "security_system": {
                    "access_control": list of strings or null,
                    "cctv": list of strings or null,
                    "intrusion_detection": list of strings or null,
                    "intercom": list of strings or null,
                    "security_zones": list of strings or null,
                    "network_infrastructure": list of strings or null
                },
                "building_infrastructure": {
                    "rooms": list of strings or null,
                    "ceiling_info": list of strings or null,
                    "wall_info": list of strings or null,
                    "electrical_rooms": list of strings or null,
                    "network_closets": list of strings or null,
                    "cable_pathways": list of strings or null
                },
                "compliance": {
                    "nfpa_codes": list of strings or null,
                    "ul_listings": list of strings or null,
                    "ada_compliance": list of strings or null,
                    "local_codes": list of strings or null,
                    "egress_paths": list of strings or null
                },
                "technical_specs": {
                    "mounting_heights": list of strings or null,
                    "coverage_areas": list of strings or null,
                    "cable_specs": list of strings or null,
                    "power_requirements": list of strings or null,
                    "network_requirements": list of strings or null
                },
                "integration_points": {
                    "building_management": list of strings or null,
                    "elevator_controls": list of strings or null,
                    "hvac_systems": list of strings or null,
                    "door_hardware": list of strings or null,
                    "emergency_power": list of strings or null
                }
            }
            
            If any information is not found, use null for that field."""),
            ("user", [
                {
                    "type": "text",
                    "text": "Please analyze this technical drawing and extract all relevant information for fire alarm and security system estimation."
                },
                {
                    "type": "image_url",
                    "image_url": {
                        "url": "{image_url}"
                    }
                }
            ])
        ])

    def _encode_image(self, image_path: Union[str, Path]) -> str:
        """Encode image to base64 for OpenAI API"""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def process_file(self, file_path: Union[str, Path]) -> List[Document]:
        """Process image file using OpenAI's Vision API"""
        try:
            # Encode image to base64
            base64_image = self._encode_image(file_path)
            image_url = f"data:image/jpeg;base64,{base64_image}"
            
            # Create the chain
            chain = self.vision_prompt | self.vision_model
            
            # Get the response
            response = chain.invoke({"image_url": image_url})
            
            # Parse the JSON response
            analysis = json.loads(response.content)
            
            # Create a formatted text representation
            content = f"""
            Technical Drawing Analysis:
            
            Document Information:
            Type: {analysis['document_info'].get('type', 'Unknown')}
            Scale: {analysis['document_info'].get('scale', 'Not specified')}
            Version: {analysis['document_info'].get('version', 'Not specified')}
            Date: {analysis['document_info'].get('date', 'Not specified')}
            
            Fire Alarm System:
            Control Panels:
            {chr(10).join(analysis['fire_alarm_system'].get('control_panels', ['None found']))}
            
            Initiating Devices:
            {chr(10).join(analysis['fire_alarm_system'].get('initiating_devices', ['None found']))}
            
            Notification Appliances:
            {chr(10).join(analysis['fire_alarm_system'].get('notification_appliances', ['None found']))}
            
            Circuits:
            {chr(10).join(analysis['fire_alarm_system'].get('circuits', ['None found']))}
            
            Power Supplies:
            {chr(10).join(analysis['fire_alarm_system'].get('power_supplies', ['None found']))}
            
            Emergency Communication:
            {chr(10).join(analysis['fire_alarm_system'].get('emergency_communication', ['None found']))}
            
            NFPA Compliance:
            {chr(10).join(analysis['fire_alarm_system'].get('nfpa_compliance', ['None found']))}
            
            Security System:
            Access Control:
            {chr(10).join(analysis['security_system'].get('access_control', ['None found']))}
            
            CCTV:
            {chr(10).join(analysis['security_system'].get('cctv', ['None found']))}
            
            Intrusion Detection:
            {chr(10).join(analysis['security_system'].get('intrusion_detection', ['None found']))}
            
            Intercom:
            {chr(10).join(analysis['security_system'].get('intercom', ['None found']))}
            
            Security Zones:
            {chr(10).join(analysis['security_system'].get('security_zones', ['None found']))}
            
            Network Infrastructure:
            {chr(10).join(analysis['security_system'].get('network_infrastructure', ['None found']))}
            
            Building Infrastructure:
            Rooms:
            {chr(10).join(analysis['building_infrastructure'].get('rooms', ['None found']))}
            
            Ceiling Information:
            {chr(10).join(analysis['building_infrastructure'].get('ceiling_info', ['None found']))}
            
            Wall Information:
            {chr(10).join(analysis['building_infrastructure'].get('wall_info', ['None found']))}
            
            Electrical Rooms:
            {chr(10).join(analysis['building_infrastructure'].get('electrical_rooms', ['None found']))}
            
            Network Closets:
            {chr(10).join(analysis['building_infrastructure'].get('network_closets', ['None found']))}
            
            Cable Pathways:
            {chr(10).join(analysis['building_infrastructure'].get('cable_pathways', ['None found']))}
            
            Compliance:
            NFPA Codes:
            {chr(10).join(analysis['compliance'].get('nfpa_codes', ['None found']))}
            
            UL Listings:
            {chr(10).join(analysis['compliance'].get('ul_listings', ['None found']))}
            
            ADA Compliance:
            {chr(10).join(analysis['compliance'].get('ada_compliance', ['None found']))}
            
            Local Codes:
            {chr(10).join(analysis['compliance'].get('local_codes', ['None found']))}
            
            Egress Paths:
            {chr(10).join(analysis['compliance'].get('egress_paths', ['None found']))}
            
            Technical Specifications:
            Mounting Heights:
            {chr(10).join(analysis['technical_specs'].get('mounting_heights', ['None found']))}
            
            Coverage Areas:
            {chr(10).join(analysis['technical_specs'].get('coverage_areas', ['None found']))}
            
            Cable Specifications:
            {chr(10).join(analysis['technical_specs'].get('cable_specs', ['None found']))}
            
            Power Requirements:
            {chr(10).join(analysis['technical_specs'].get('power_requirements', ['None found']))}
            
            Network Requirements:
            {chr(10).join(analysis['technical_specs'].get('network_requirements', ['None found']))}
            
            Integration Points:
            Building Management:
            {chr(10).join(analysis['integration_points'].get('building_management', ['None found']))}
            
            Elevator Controls:
            {chr(10).join(analysis['integration_points'].get('elevator_controls', ['None found']))}
            
            HVAC Systems:
            {chr(10).join(analysis['integration_points'].get('hvac_systems', ['None found']))}
            
            Door Hardware:
            {chr(10).join(analysis['integration_points'].get('door_hardware', ['None found']))}
            
            Emergency Power:
            {chr(10).join(analysis['integration_points'].get('emergency_power', ['None found']))}
            """
            
            return self.process(content)
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            return self.process("")

class CADProcessor(DocumentProcessor):
    """Process CAD files (DXF)"""
    def process_file(self, file_path: Union[str, Path]) -> List[Document]:
        try:
            doc = ezdxf.readfile(str(file_path))
            
            # Extract layers
            layers = [layer.dxf.name for layer in doc.layers]
            
            # Extract entities
            entities = []
            for entity in doc.modelspace():
                entities.append(f"{entity.dxftype()}: {entity.dxf.handle}")
            
            # Extract blocks
            blocks = [block.name for block in doc.blocks]
            
            content = f"""
            CAD File Analysis:
            - Layers: {', '.join(layers)}
            - Entities: {len(entities)}
            - Blocks: {', '.join(blocks)}
            """
            
            return self.process(content)
        except Exception as e:
            logger.error(f"Error processing CAD file: {e}")
            return self.process("")

class EmailProcessor:
    """Process emails using Gmail API or AWS SES"""
    def __init__(self, service_type: str = "gmail", redis_url: Optional[str] = None):
        self.service_type = service_type
        self.processor = DocumentProcessor(redis_url)
        if service_type == "gmail":
            self.service = self._setup_gmail_service()
        else:
            self.ses_client = boto3.client('ses')

    def _setup_gmail_service(self):
        """Setup Gmail API service"""
        SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
        creds = None
        
        if os.path.exists('token.pickle'):
            with open('token.pickle', 'rb') as token:
                creds = pickle.load(token)
        
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(
                    'credentials.json', SCOPES)
                creds = flow.run_local_server(port=0)
            
            with open('token.pickle', 'wb') as token:
                pickle.dump(creds, token)

        return build('gmail', 'v1', credentials=creds)

    def process_email(self, email_id: str) -> List[Document]:
        """Process email content"""
        if self.service_type == "gmail":
            message = self.service.users().messages().get(
                userId='me', id=email_id, format='full').execute()
            
            # Extract email body
            if 'payload' in message and 'parts' in message['payload']:
                parts = message['payload']['parts']
                content = ""
                for part in parts:
                    if part['mimeType'] == 'text/plain':
                        data = part['body'].get('data', '')
                        content += base64.urlsafe_b64decode(data).decode()
            else:
                data = message['payload']['body'].get('data', '')
                content = base64.urlsafe_b64decode(data).decode()
        else:
            # AWS SES implementation
            response = self.ses_client.get_email_identity(
                EmailAddress=email_id
            )
            content = response['EmailIdentity']

        return self.processor.process(content)

class DocumentIngestionPipeline:
    """Main pipeline for document ingestion"""
    def __init__(self, redis_url: Optional[str] = None):
        self.processors = {
            'application/pdf': PDFProcessor(redis_url),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocxProcessor(redis_url),
            'text/plain': DocumentProcessor(redis_url),
            'image/jpeg': ImageProcessor(redis_url),
            'image/png': ImageProcessor(redis_url),
            'image/gif': ImageProcessor(redis_url),
            'image/bmp': ImageProcessor(redis_url),
            'application/dxf': CADProcessor(redis_url),
        }

    def process_file(self, file_path: Union[str, Path]) -> List[Document]:
        """Process a file based on its MIME type"""
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(str(file_path))
        
        if file_type in self.processors:
            return self.processors[file_type].process_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def process_email(self, email_id: str, service_type: str = "gmail") -> List[Document]:
        """Process an email"""
        processor = EmailProcessor(service_type)
        return processor.process_email(email_id)

    def process_directory(self, directory_path: Union[str, Path]) -> List[Document]:
        """Process all supported files in a directory"""
        all_documents = []
        for file_path in Path(directory_path).rglob("*"):
            if file_path.is_file():
                try:
                    documents = self.process_file(file_path)
                    all_documents.extend(documents)
                except ValueError as e:
                    logger.warning(f"Skipping unsupported file {file_path}: {e}")
                    continue
                except Exception as e:
                    logger.error(f"Error processing file {file_path}: {e}")
                    continue
        return all_documents 